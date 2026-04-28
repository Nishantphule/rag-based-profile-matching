"""
resume_rag.py
=============

Part A: RAG System Setup.

Implements a complete document-processing pipeline for resumes:

1. Document loading (.txt, .pdf, .docx)
2. Section-aware chunking (Summary, Skills, Experience, Education, Projects, ...)
3. Metadata extraction (name, skills, years of experience, education)
4. Embedding via SentenceTransformers (all-MiniLM-L6-v2)
5. Persistent storage in ChromaDB

Usage:
    # Build / refresh the index
    python resume_rag.py --build --resumes-dir data/resumes

    # Show stats
    python resume_rag.py --stats

The same classes are imported by `job_matcher.py` and the Jupyter notebook.
"""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import time
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Iterable

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
from tqdm import tqdm

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(name)s | %(message)s",
)
# Quiet third-party noise so the demo output stays clean.
for _noisy in ("httpx", "httpcore", "huggingface_hub", "sentence_transformers",
               "urllib3", "chromadb", "openai"):
    logging.getLogger(_noisy).setLevel(logging.ERROR)
log = logging.getLogger("resume_rag")


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

DEFAULT_EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_COLLECTION = "resumes"
DEFAULT_DB_PATH = "chroma_db"

# Canonical section names. Anything we detect maps to one of these labels.
SECTION_LABELS = [
    "summary",
    "skills",
    "experience",
    "education",
    "projects",
    "certifications",
    "awards",
    "publications",
    "header",   # name + contact details at top
    "other",
]

# Heading aliases -> canonical label.
SECTION_ALIASES = {
    "summary": ["summary", "profile", "objective", "about me", "professional summary"],
    "skills": ["skills", "technical skills", "core competencies", "tech stack", "technologies"],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history",
    ],
    "education": ["education", "academics", "academic background", "qualifications"],
    "projects": ["projects", "personal projects", "side projects", "selected projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "awards": ["awards", "honors", "achievements"],
    "publications": ["publications", "papers"],
}

# Reverse lookup: lowercase heading text -> canonical label.
HEADING_TO_LABEL: dict[str, str] = {}
for label, aliases in SECTION_ALIASES.items():
    for alias in aliases:
        HEADING_TO_LABEL[alias] = label


# A reasonably broad skill vocabulary used for metadata extraction and BM25.
# It is not meant to be exhaustive - it's a guard-rail to keep extraction
# deterministic and cheap.
SKILL_VOCAB: list[str] = sorted(
    {
        # Languages
        "python", "java", "scala", "kotlin", "swift", "objective-c", "go", "rust",
        "c", "c++", "javascript", "typescript", "r", "sql", "bash", "shell",
        # Web / frontend
        "react", "next.js", "node.js", "express", "django", "flask", "fastapi",
        "spring boot", "tailwind css", "storybook", "webpack",
        # Data / ML / AI
        "machine learning", "deep learning", "nlp", "computer vision", "mlops",
        "pytorch", "tensorflow", "jax", "scikit-learn", "pandas", "numpy",
        "transformers", "hugging face", "langchain", "llamaindex", "llms",
        "rag", "prompt engineering", "openai", "cohere",
        "vector databases", "chromadb", "pinecone", "weaviate", "faiss",
        "mlflow", "airflow", "kubeflow", "spark", "spark streaming", "flink",
        "kafka", "snowflake", "databricks", "dbt",
        # Infra / cloud
        "aws", "gcp", "azure", "kubernetes", "docker", "terraform", "helm",
        "argocd", "prometheus", "grafana", "linux", "redis", "postgresql",
        "mysql", "mongodb",
        # CV / robotics
        "opencv", "ros", "ros2", "tensorrt", "cuda",
        # Mobile
        "android", "ios", "jetpack compose", "swiftui", "combine",
        # Misc
        "system design", "microservices", "a/b testing", "causal inference",
        "statistics", "tableau", "power bi", "looker",
        "selenium", "playwright", "cypress", "jmeter", "appium",
    }
)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class ResumeChunk:
    """A single retrievable chunk extracted from a resume."""

    chunk_id: str
    resume_id: str
    section: str          # canonical section label
    text: str
    metadata: dict = field(default_factory=dict)


@dataclass
class ResumeDocument:
    """Parsed representation of a single resume on disk."""

    resume_id: str
    file_path: str
    raw_text: str
    sections: dict[str, str] = field(default_factory=dict)   # canonical label -> text
    metadata: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# 1. Loading
# ---------------------------------------------------------------------------

class ResumeLoader:
    """Load resumes from disk. Supports .txt, .pdf and .docx."""

    SUPPORTED_EXT = (".txt", ".md", ".pdf", ".docx")

    def load_dir(self, directory: str | Path) -> list[ResumeDocument]:
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Resume directory not found: {directory}")

        files = [p for p in sorted(directory.iterdir()) if p.suffix.lower() in self.SUPPORTED_EXT]
        log.info("Found %d resume files in %s", len(files), directory)

        out: list[ResumeDocument] = []
        for fp in files:
            try:
                out.append(self.load_file(fp))
            except Exception as exc:  # pragma: no cover - defensive
                log.warning("Failed to load %s: %s", fp, exc)
        return out

    def load_file(self, path: str | Path) -> ResumeDocument:
        path = Path(path)
        text = self._read_text(path)
        return ResumeDocument(
            resume_id=path.stem,
            file_path=str(path),
            raw_text=text,
        )

    def _read_text(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")
        if ext == ".pdf":
            from pypdf import PdfReader  # lazy import
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if ext == ".docx":
            import docx  # python-docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError(f"Unsupported file extension: {ext}")


# ---------------------------------------------------------------------------
# 2. Section-aware chunking
# ---------------------------------------------------------------------------

class ResumeChunker:
    """Split a resume into chunks that preserve resume sections.

    Strategy:
        1. Detect section headers (e.g. "EXPERIENCE", "Skills:") and split.
        2. Each section becomes its own chunk; long sections are sub-chunked
           by entry (a bullet group) or by a soft token-budget so we stay
           under the embedding model's context window.
    """

    # Heuristic regex for headings: a short, mostly-uppercase line OR a line
    # whose normalized text matches a known alias. Keep the rule simple - we
    # validate with an alias lookup right after.
    HEADING_RE = re.compile(r"^\s*([A-Za-z][A-Za-z &/\-]{2,40})\s*[:\-]?\s*$")

    def __init__(self, max_chars: int = 1200) -> None:
        self.max_chars = max_chars

    # ---- public API ----------------------------------------------------

    def chunk(self, doc: ResumeDocument) -> list[ResumeChunk]:
        sections = self._split_sections(doc.raw_text)
        doc.sections = sections

        chunks: list[ResumeChunk] = []
        for label, text in sections.items():
            text = text.strip()
            if not text:
                continue
            for i, sub in enumerate(self._subchunk(text, label)):
                chunks.append(
                    ResumeChunk(
                        chunk_id=f"{doc.resume_id}::{label}::{i}",
                        resume_id=doc.resume_id,
                        section=label,
                        text=sub,
                    )
                )
        return chunks

    # ---- internals -----------------------------------------------------

    def _split_sections(self, text: str) -> dict[str, str]:
        """Walk lines and split on detected headings."""
        lines = text.splitlines()
        sections: dict[str, list[str]] = {"header": []}
        current = "header"

        for raw_line in lines:
            label = self._detect_heading(raw_line)
            if label is not None:
                current = label
                sections.setdefault(current, [])
                continue
            sections.setdefault(current, []).append(raw_line)

        # Join, strip, drop empties.
        return {k: "\n".join(v).strip() for k, v in sections.items() if "".join(v).strip()}

    def _detect_heading(self, line: str) -> str | None:
        """Return canonical label if `line` looks like a section heading."""
        stripped = line.strip().rstrip(":").strip()
        if not stripped or len(stripped) > 50:
            return None

        normalized = re.sub(r"\s+", " ", stripped.lower())
        if normalized in HEADING_TO_LABEL:
            # Visual filters: heading lines are usually short and either fully
            # title-case / upper-case, OR end with a colon. Accept either.
            is_titleish = stripped.isupper() or stripped.istitle() or line.rstrip().endswith(":")
            if is_titleish or len(stripped.split()) <= 3:
                return HEADING_TO_LABEL[normalized]
        return None

    def _subchunk(self, text: str, label: str) -> list[str]:
        """Chunk a section's text, respecting bullets and a char budget."""
        if len(text) <= self.max_chars:
            return [text]

        # Try to split on top-level bullet entries first (lines starting with "- ").
        bullet_groups: list[list[str]] = []
        current: list[str] = []
        for line in text.splitlines():
            if re.match(r"^\s*[-*•]\s+", line) and current:
                bullet_groups.append(current)
                current = [line]
            else:
                current.append(line)
        if current:
            bullet_groups.append(current)

        # Greedy pack groups into chunks under max_chars.
        chunks: list[str] = []
        buf: list[str] = []
        buf_len = 0
        for group in bullet_groups:
            group_text = "\n".join(group).strip()
            if not group_text:
                continue
            if buf_len + len(group_text) + 1 > self.max_chars and buf:
                chunks.append("\n".join(buf).strip())
                buf, buf_len = [], 0
            buf.append(group_text)
            buf_len += len(group_text) + 1
        if buf:
            chunks.append("\n".join(buf).strip())

        # Safety net: if a single bullet was huge, hard-wrap by character.
        out: list[str] = []
        for c in chunks:
            if len(c) <= self.max_chars:
                out.append(c)
            else:
                for i in range(0, len(c), self.max_chars):
                    out.append(c[i : i + self.max_chars])
        return out


# ---------------------------------------------------------------------------
# 3. Metadata extraction
# ---------------------------------------------------------------------------

class MetadataExtractor:
    """Extract Name, Skills, Years of Experience, Education from a parsed resume."""

    EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
    PHONE_RE = re.compile(r"(\+?\d[\d\-\s]{7,}\d)")

    YEARS_RES = [
        re.compile(r"total\s+experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*years?", re.I),
        re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+of\s+experience", re.I),
        re.compile(r"experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*years?", re.I),
    ]

    DEGREE_RES = re.compile(
        r"\b(?:Ph\.?D|M\.?Tech|M\.?Sc|MS|MBA|M\.?E|M\.?S|"
        r"B\.?Tech|B\.?Sc|BS|BA|B\.?E|B\.?Com|B\.?Des)\b[^,\n]*",
        re.I,
    )

    def extract(self, doc: ResumeDocument) -> dict:
        """Populate and return `doc.metadata`."""
        text = doc.raw_text
        meta: dict = {
            "resume_id": doc.resume_id,
            "file_path": doc.file_path,
            "name": self._extract_name(doc),
            "email": self._first(self.EMAIL_RE, text),
            "phone": self._first(self.PHONE_RE, text),
            "experience_years": self._extract_years(text),
            "skills": self._extract_skills(text),
            "education": self._extract_education(doc),
            "title": self._extract_title(doc),
        }
        # Chroma metadata can only contain primitives, so we serialize lists.
        meta["skills_str"] = ", ".join(meta["skills"])
        meta["education_str"] = " | ".join(meta["education"]) if meta["education"] else ""
        doc.metadata = meta
        return meta

    # ---- field-by-field helpers ---------------------------------------

    @staticmethod
    def _first(pattern: re.Pattern[str], text: str) -> str:
        m = pattern.search(text)
        return m.group(0).strip() if m else ""

    def _extract_name(self, doc: ResumeDocument) -> str:
        # Strategy: first non-empty line of the header section, capped to ~5
        # words and not containing an email/phone.
        header = doc.sections.get("header") or doc.raw_text
        for line in header.splitlines():
            line = line.strip()
            if not line:
                continue
            if self.EMAIL_RE.search(line) or self.PHONE_RE.search(line):
                continue
            words = line.split()
            if 1 <= len(words) <= 6 and all(w[0].isupper() or "." in w for w in words if w):
                return line
        # Fallback: filename
        return doc.resume_id.replace("_", " ").title()

    def _extract_title(self, doc: ResumeDocument) -> str:
        header = doc.sections.get("header", "")
        lines = [l.strip() for l in header.splitlines() if l.strip()]
        # Title is usually the 2nd non-empty header line.
        return lines[1] if len(lines) > 1 else ""

    def _extract_years(self, text: str) -> float:
        for r in self.YEARS_RES:
            m = r.search(text)
            if m:
                try:
                    return float(m.group(1))
                except ValueError:
                    continue
        return 0.0

    def _extract_skills(self, text: str) -> list[str]:
        text_lc = text.lower()
        found: list[str] = []
        for skill in SKILL_VOCAB:
            # word-boundary match for short tokens, substring for multi-word
            if " " in skill or "." in skill or "+" in skill or "-" in skill:
                if skill in text_lc:
                    found.append(skill)
            else:
                if re.search(rf"\b{re.escape(skill)}\b", text_lc):
                    found.append(skill)
        # Stable, dedup while preserving canonical order
        return sorted(set(found))

    def _extract_education(self, doc: ResumeDocument) -> list[str]:
        edu_text = doc.sections.get("education", "")
        if not edu_text:
            edu_text = doc.raw_text
        matches = self.DEGREE_RES.findall(edu_text)
        # Trim and dedup, preserving order.
        seen: set[str] = set()
        out: list[str] = []
        for m in matches:
            cleaned = re.sub(r"\s+", " ", m).strip(" ,.;:")
            if cleaned and cleaned.lower() not in seen:
                seen.add(cleaned.lower())
                out.append(cleaned)
        return out


# ---------------------------------------------------------------------------
# 4. Embedding + Vector Store (ChromaDB)
# ---------------------------------------------------------------------------

class Embedder:
    """Thin wrapper around SentenceTransformer with batching."""

    def __init__(self, model_name: str = DEFAULT_EMBED_MODEL) -> None:
        log.info("Loading embedding model: %s", model_name)
        self.model_name = model_name
        self.model = SentenceTransformer(model_name)

    def encode(self, texts: list[str], batch_size: int = 64) -> list[list[float]]:
        if not texts:
            return []
        vecs = self.model.encode(
            texts,
            batch_size=batch_size,
            convert_to_numpy=True,
            normalize_embeddings=True,
            show_progress_bar=False,
        )
        return vecs.tolist()


class ResumeRAG:
    """Top-level orchestrator: load -> chunk -> extract -> embed -> store."""

    def __init__(
        self,
        db_path: str = DEFAULT_DB_PATH,
        collection_name: str = DEFAULT_COLLECTION,
        embed_model: str = DEFAULT_EMBED_MODEL,
    ) -> None:
        self.db_path = db_path
        self.collection_name = collection_name

        self.loader = ResumeLoader()
        self.chunker = ResumeChunker()
        self.extractor = MetadataExtractor()
        self.embedder = Embedder(embed_model)

        self.client = chromadb.PersistentClient(
            path=db_path,
            settings=Settings(anonymized_telemetry=False),
        )
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"hnsw:space": "cosine"},
        )

    # ---- ingestion -----------------------------------------------------

    def build(self, resumes_dir: str | Path, reset: bool = True) -> dict:
        """Load resumes and (re)build the vector index. Returns build stats."""
        if reset:
            log.info("Resetting collection '%s'", self.collection_name)
            try:
                self.client.delete_collection(self.collection_name)
            except Exception:
                pass
            self.collection = self.client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"},
            )

        t0 = time.perf_counter()
        docs = self.loader.load_dir(resumes_dir)

        all_chunks: list[ResumeChunk] = []
        for d in tqdm(docs, desc="Parsing resumes"):
            chunks = self.chunker.chunk(d)         # also populates d.sections
            self.extractor.extract(d)              # populates d.metadata
            for chunk in chunks:
                chunk.metadata = {
                    **{k: v for k, v in d.metadata.items() if not isinstance(v, (list, dict))},
                    "section": chunk.section,
                }
                all_chunks.append(chunk)

        log.info("Embedding %d chunks from %d resumes ...", len(all_chunks), len(docs))
        texts = [c.text for c in all_chunks]
        embeddings = self.embedder.encode(texts)

        # ChromaDB ingest
        self.collection.add(
            ids=[c.chunk_id for c in all_chunks],
            documents=texts,
            embeddings=embeddings,
            metadatas=[c.metadata for c in all_chunks],
        )

        # Save resume-level metadata index for downstream filtering.
        meta_path = Path(self.db_path) / "resume_metadata.json"
        meta_path.parent.mkdir(parents=True, exist_ok=True)
        meta_path.write_text(
            json.dumps([d.metadata for d in docs], indent=2, ensure_ascii=False),
            encoding="utf-8",
        )

        elapsed = time.perf_counter() - t0
        stats = {
            "resumes": len(docs),
            "chunks": len(all_chunks),
            "elapsed_seconds": round(elapsed, 2),
            "avg_chunks_per_resume": round(len(all_chunks) / max(1, len(docs)), 2),
            "collection": self.collection_name,
            "db_path": self.db_path,
        }
        log.info("Build complete: %s", stats)
        return stats

    # ---- read APIs ----------------------------------------------------

    def list_resumes(self) -> list[dict]:
        meta_path = Path(self.db_path) / "resume_metadata.json"
        if not meta_path.exists():
            return []
        return json.loads(meta_path.read_text(encoding="utf-8"))

    def stats(self) -> dict:
        return {
            "collection": self.collection_name,
            "num_chunks": self.collection.count(),
            "num_resumes": len(self.list_resumes()),
            "db_path": self.db_path,
        }

    def query(
        self,
        text: str,
        top_k: int = 10,
        where: dict | None = None,
    ) -> dict:
        """Raw semantic query exposed for the matcher and the notebook."""
        emb = self.embedder.encode([text])[0]
        return self.collection.query(
            query_embeddings=[emb],
            n_results=top_k,
            where=where,
            include=["documents", "metadatas", "distances"],
        )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Resume RAG indexer (Part A)")
    p.add_argument("--build", action="store_true", help="Build/refresh the vector index")
    p.add_argument("--stats", action="store_true", help="Print collection stats")
    p.add_argument("--resumes-dir", default="data/resumes", help="Directory of resumes")
    p.add_argument("--db-path", default=DEFAULT_DB_PATH)
    p.add_argument("--collection", default=DEFAULT_COLLECTION)
    p.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    p.add_argument("--no-reset", action="store_true", help="Do not delete the existing collection")
    p.add_argument("--json", action="store_true",
                   help="Print stats as raw JSON instead of the rich summary panel.")
    return p.parse_args()


def main() -> None:
    args = _parse_args()

    # Optional rich renderer. When active, silence our own INFO logs so the
    # rich panels render cleanly (warnings/errors still come through).
    pretty = None
    if not args.json:
        try:
            from cli_pretty import banner, render_index_stats
            pretty = render_index_stats
            log.setLevel(logging.WARNING)
            banner(subtitle="Indexing resumes for retrieval...")
        except Exception:  # pragma: no cover
            pretty = None

    rag = ResumeRAG(
        db_path=args.db_path,
        collection_name=args.collection,
        embed_model=args.embed_model,
    )

    if args.build:
        stats = rag.build(args.resumes_dir, reset=not args.no_reset)
        if pretty:
            pretty(stats)
        else:
            print(json.dumps(stats, indent=2))

    if args.stats or not args.build:
        s = rag.stats()
        if pretty:
            pretty({**s, "elapsed_seconds": "-", "avg_chunks_per_resume": "-",
                    "resumes": s.get("num_resumes", 0), "chunks": s.get("num_chunks", 0)})
        else:
            print(json.dumps(s, indent=2))


if __name__ == "__main__":
    main()
